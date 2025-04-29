import marimo

__generated_with = "0.3.2"
app = marimo.App(width="full")


@app.cell
def __():
    from collections import defaultdict
    from plotly.subplots import make_subplots
    import plotly.graph_objects as go
    return defaultdict, go, make_subplots


@app.cell
def __():
    import marimo as mo
    import pandas as pd
    import jsonlines
    import numpy as np

    import plotly.express as px


    from vdl_tools.shared_tools.project_relevance.add_categorical_skills_to_linkedin_3 import (
        calculate_project_relevance,
        set_project_relevance,
        add_keywords,
        PROJECT_RELEVANCE_THRESHOLD,
        PROJECT_RELEVANCE_IDF_THRESHOLD,
        TECH_RELEVANCE_THRESHOLD,
    )
    from vdl_tools.shared_tools.project_relevance.add_founder_toplines_to_venture_cft_4 import (
        add_project_relevance_to_venture_df,
    )

    from vdl_tools.shared_tools.project_config import get_paths


    PATHS = get_paths()
    return (
        PATHS,
        PROJECT_RELEVANCE_IDF_THRESHOLD,
        PROJECT_RELEVANCE_THRESHOLD,
        TECH_RELEVANCE_THRESHOLD,
        add_keywords,
        add_project_relevance_to_venture_df,
        calculate_project_relevance,
        get_paths,
        jsonlines,
        mo,
        np,
        pd,
        px,
        set_project_relevance,
    )


@app.cell
def __(PATHS, jsonlines):
    with jsonlines.open(PATHS["sensitivity_analysis_founders"], "r") as _f:
        simulations_founders = [x for x in _f]
    return simulations_founders,


@app.cell
def __(defaultdict, simulations_founders):
    founder_li_id_to_founder_sim_list = defaultdict(list)

    for founder_list in simulations_founders:
        for founder in founder_list:
            founder_li_id_to_founder_sim_list[founder["li_id"]].append(founder)
    return founder, founder_li_id_to_founder_sim_list, founder_list


@app.cell
def __(pd, simulations_founders):
    def get_master_df():
        def _get_run_df():
            for founder_list in simulations_founders:
                yield pd.DataFrame(founder_list)

        return pd.concat(_get_run_df(), ignore_index=True)
    return get_master_df,


@app.cell
def __(get_master_df):
    master_df = get_master_df()
    return master_df,


@app.cell
def __(master_df, np):
    def calc_25(x):
        return np.percentile(x, 25)


    def calc_75(x):
        return np.percentile(x, 75)


    metrics_df = master_df.groupby("li_id")[
        [
            "tech_relevant_ratio",
            "project_relevant_ratio",
            "project_relevant_idf_ratio",
        ]
    ].agg(
        [
            np.mean,
            np.std,
            np.median,
            calc_25,
            calc_75,
        ]
    )
    return calc_25, calc_75, metrics_df


@app.cell
def __(go, make_subplots, np, px):
    def create_bar_neg_cum_sum(
        metrics_df,
        column,
        metric,
        n_bins=50,
        normalize=False,
        title=None,
        threshold=None,
        color="limegreen",
    ):
        data = metrics_df[column][metric]

        n_bins = n_bins + 1
        step = (max(data) - min(data)) / n_bins
        bins = np.arange(min(data), max(data), step)

        counts, bins = np.histogram(data, bins=bins)
        bins = 0.5 * (bins[:-1] + bins[1:])

        cum_sum = [sum(counts[:i]) for i in range(len(counts))]

        if normalize:
            y = (max(cum_sum) - cum_sum) / max(cum_sum)
        else:
            y = max(cum_sum) - cum_sum
        fig = px.bar(x=bins, y=y, color_discrete_sequence=[color])

        title = title or f"{column} {metric}"
        fig.update_layout(
            title=title,
            xaxis_title=column,
            yaxis_title="Cumulative Count",
        )
        if threshold:
            fig.add_vline(x=threshold, line_dash="dash", line_color="blue")
        return fig


    def create_side_by_side(
        metrics_df,
        column,
        metric,
        threshold=None,
        title=None,
    ):

        fig = go.Figure()
        _fig_abs = create_bar_neg_cum_sum(
            metrics_df,
            column,
            metric,
            threshold=threshold,
        )
        _fig_norm = create_bar_neg_cum_sum(
            metrics_df,
            column,
            metric,
            normalize=True,
            threshold=threshold,
        )
        fig = make_subplots(
            rows=1, cols=2, subplot_titles=("Absolute", "Normalized")
        )
        fig.add_trace(_fig_abs["data"][0], row=1, col=1)
        fig.add_trace(_fig_norm["data"][0], row=1, col=2)
        if threshold:
            fig.add_vline(
                x=threshold, line_dash="dash", line_color="blue", row=1, col=1
            )
            fig.add_vline(
                x=threshold, line_dash="dash", line_color="blue", row=1, col=2
            )

        title = title or f"{column} {metric}"
        fig.update_layout(
            title=title,
            yaxis_title="Cumulative count at threshold",
            xaxis_title="Ratio",
        )
        return fig
    return create_bar_neg_cum_sum, create_side_by_side


@app.cell
def __(mo):
    mo.md(
        f"""
    # Founders w/Project Relevance

    ## Recap of Previous Work

    - We have calculated the project relevance of each founder by looking at the skills they have listed on their LinkedIn profile.
    - Labeled a set of skills to be project relevant or not.
    - Calculated the ratio of project relevant skills to total skills for each founder
    - If the founder had more than 10% of keywords in their **skills**, **experiences**, or **summary** as "project relevant", they were considered to have "project relevant experience"
    - This resulted in **~20% of founders** to have project relevant experience and **~30% of companies had at least 1 founder** with project relevant experience

    ### Sensitivity Analysis
    - Because the 10% threshold and selected skills were arbitrary, we performed a sensitivity analysis to determine the impact of these choices on the results.
     - We varied the threshold by a normal distribution centered on 10% with a standard deviation of 5%
     - We also randomly flipped the "project relevant" label on with a .5% probablity for each term
     - We randomly deleted / added skills for each founder
     - Ran 1000 times
     - Again, **~30% of companies** had at least 1 founder with project relevant experience with the IQR being between ~15% and 50%

    ## Updates from Amir
    - Amir at Elemental gave feedback on the "project relevant" terms and suggested some changes. Those changes brought in some terms that were much more common, like:
    'entrepreneurship' and 'strategy'

    - Because of this, the ratio of project relevant terms starts to lose signal since these terms cover a lot of people
    - Added a weight of the IDF of the term to account for common terms

    ## Next Steps
    - Using the IDF scores, we need to come up with a new threshold to center the sensitivity analysis around  
    - **But how to choose a sane one? I eyeballed it so it would include about 20% of founders as having the experience?**
    - Maggie's colleague also gave feedback and is a lot stricter than Amir and contradicts some of those.  [See here](https://docs.google.com/spreadsheets/d/1Cnwpdr3RvXFYU5ViRVtD-Z-LBD_U0u_ZWvc6G1374Cc/edit#gid=0)
    - **How to reconcile?**
    - **How to get feedback that the threshold is sane?**
    - **Idea** Create profiles of the "highest scoring", "lowest scoring", and some around the threshold. Show those to Amir, Lee, and get feedback
    """
    )
    return


@app.cell
def __(mo):
    mo.md("""### Distribution of Founders "Project Relevant" Ratios""")
    return


@app.cell
def __(PROJECT_RELEVANCE_THRESHOLD, create_side_by_side, metrics_df):
    create_side_by_side(
        metrics_df,
        "project_relevant_ratio",
        "mean",
        threshold=PROJECT_RELEVANCE_THRESHOLD,
    )
    return


@app.cell
def __(PROJECT_RELEVANCE_IDF_THRESHOLD, create_side_by_side, metrics_df):
    create_side_by_side(
        metrics_df,
        "project_relevant_idf_ratio",
        "mean",
        threshold=PROJECT_RELEVANCE_IDF_THRESHOLD,
    )
    return


@app.cell
def __(mo):
    mo.md("## Appendix")
    return


@app.cell
def __(PATHS, get_experience_string, metrics_df, pd):
    founders_df_raw = pd.read_json(PATHS["annotated_linkedin_profiles"])

    drop_columns = [
        "middle_skill_counts",
        "parent_skill_counts",
        "project_relevant_sum",
        "project_relevant_idf_sum",
        "tech_relevant_sum",
        "tech_large_sum",
        "tech_relevant_idf_sum",
        "tech_large_idf_sum",
        "tech_relevant_ratio",
        "project_relevant_ratio",
        "tech_large_ratio",
        "tech_relevant_idf_ratio",
        "project_relevant_idf_ratio",
        "tech_large_idf_ratio",
        "is_project_relevant",
        "is_project_relevant_idf",
        "is_tech_majority",
        "is_tech_large_majority",
    ]

    founders_df = founders_df_raw.copy()
    founders_df = founders_df.drop(columns=drop_columns, inplace=False)

    _metrics_df = metrics_df.reset_index()
    founders_df = founders_df.merge(
        _metrics_df, left_on="li_id", right_on="li_id", how="left"
    ).copy()

    founders_df["experience_full"] = founders_df[
        "member_experience_collection"
    ].apply(get_experience_string)
    return drop_columns, founders_df, founders_df_raw


@app.cell
def __(founders_df):
    key_columns = [
        "name",
        "li_id",
        "url",
        "skills",
        "skills_tags_list",
        "experience_full",
        "summary",
        "matched_project_relevant_terms",
        ("project_relevant_ratio", "median"),
        ("project_relevant_idf_ratio", "median"),
    ]
    table_df = founders_df[key_columns].copy()
    table_df["project_relevant_ratio"] = table_df.pop(
        ("project_relevant_ratio", "median")
    )
    table_df["project_relevant_idf_ratio"] = table_df.pop(
        ("project_relevant_idf_ratio", "median")
    )

    table_df.sort_values(
        "project_relevant_idf_ratio", ascending=False, inplace=True
    )
    return key_columns, table_df


@app.cell
def __(mo, table_df):
    _table = mo.ui.table(table_df[:100], page_size=10)

    mo.md(f"## Top Profiles{_table}")
    return


@app.cell
def __(mo, table_df):
    _table = mo.ui.table(
        table_df[table_df["project_relevant_idf_ratio"].between(0.68, 0.71)][:100],
        page_size=10,
    )
    mo.md(
        f"""## Mid Profiles
    {_table}
    """
    )
    return


@app.cell
def __(mo, table_df):
    _table = mo.ui.table(
        table_df[table_df["project_relevant_idf_ratio"].between(0.3, 0.4)][:100],
        page_size=10,
    )
    mo.md(
        f"""## Lower Profiles
    {_table}
    """
    )
    return


@app.cell
def __(PROJECT_RELEVANCE_IDF_THRESHOLD, founders_df, mo, table_df):
    project_relevant_idf_relevant_mask = (
        founders_df[("project_relevant_idf_ratio", "median")]
        >= PROJECT_RELEVANCE_IDF_THRESHOLD
    )
    project_relevant_mask = (
        founders_df[("project_relevant_ratio", "median")] >= 0.2
    )

    pos_idf_neg_reg = project_relevant_idf_relevant_mask & ~project_relevant_mask
    neg_idf_pos_reg = ~project_relevant_idf_relevant_mask & project_relevant_mask

    mixed_labels_df = table_df[(pos_idf_neg_reg | neg_idf_pos_reg)]
    _table = mo.ui.table(mixed_labels_df, page_size=10)
    mo.md(
        f"""## Mixed Profiles
    {_table}
    """
    )
    return (
        mixed_labels_df,
        neg_idf_pos_reg,
        pos_idf_neg_reg,
        project_relevant_idf_relevant_mask,
        project_relevant_mask,
    )


@app.cell
def __(founders_df, px):
    fig = px.scatter(
        x=founders_df[("project_relevant_ratio", "median")],
        y=founders_df[("project_relevant_idf_ratio", "median")],
        hover_name=founders_df.apply(
            lambda x: f"{x['name']} {x['li_id']} {x['skills_tags_list']}", axis=1
        ),
        title="Project Relevance Ratios",
        labels={
            "x": "Project Relevance Ratio",
            "y": "Project Relevance IDF Ratio",
        },
        template="plotly_white",
        color_discrete_sequence=["limegreen"],
    )

    # Overlay a blue rectangle starting at x = .2 and x = 1
    fig.add_shape(
        type="rect",
        x0=0.2,
        y0=0.0,
        x1=0.7,
        y1=3.5,
        line=dict(
            color="RoyalBlue",
        ),
        fillcolor="RoyalBlue",
        opacity=0.2,
        layer="below",
    )
    fig.add_shape(
        type="rect",
        x0=0,
        y0=0.7,
        x1=0.7,
        y1=3.5,
        line=dict(
            color="red",
        ),
        fillcolor="red",
        opacity=0.2,
        layer="below",
    )
    fig
    return fig,


@app.cell
def __():
    import re


    def cleanhtml(raw_html):
        if not raw_html:
            return ""
        replaced_breaks = raw_html.replace("<br>", "\n")
        CLEANR = re.compile("<.*?>")
        cleantext = re.sub(CLEANR, "", replaced_breaks)
        return cleantext


    def format_experience(experience):
        return "\n".join(
            [
                f'dates: {experience["date_from"]} - {experience["date_to"]}',
                f'title: {experience["title"]}',
                f'company_name: {experience["company_name"]}',
                f'description:\n{cleanhtml(experience["description"])}',
            ]
        )


    def get_experience_string(experience_collection):
        return "\n-----------------\n".join(
            [
                format_experience(x)
                for x in sorted(
                    experience_collection, key=lambda x: x.get("start_year", 0), reverse=True
                )
                if x.get("title") or x.get("description")
            ]
        )
    return cleanhtml, format_experience, get_experience_string, re


@app.cell
def __(PROJECT_RELEVANCE_IDF_THRESHOLD, np, px):
    px.histogram(
        np.random.normal(PROJECT_RELEVANCE_IDF_THRESHOLD, 0.05, 1000),
        title="Distribution of Project Relevance IDF Thresholds in Simulations",
    )
    return


@app.cell
def __():
    import docx
    return docx,


@app.cell
def __():
    # def add_founder(document, founder):
    #     document.add_heading(f"{founder.name}", level=1)
    #     document.add_paragraph({founder.url})

    #     table = document.add_table(rows=1, cols=2)
    #     table.style = "Light Grid"
    #     table.style.outer_border = True
    #     row_cells = table.rows[0].cells
    #     row_cells[0].text = "is_project_relevant (please type Y / N)"
    #     row = table.add_row()
    #     row.cells[0].text = "linkedin_id"
    #     row.cells[1].text = founder.li_id
    #     document.add_paragraph()
    #     data = [
    #         (f"Skills", f"{founder.skills}"),
    #         (f"Summary", f"{founder.summary}"),
    #         (f"Experience", f"{founder.experience_full}"),
    #         (f"Skills Tags", f"{founder.skills_tags_list}"),
    #     ]

    #     for header, value in data:
    #         paragraph = document.add_paragraph()
    #         paragraph.add_run(header).bold = True
    #         paragraph.add_run("\n")
    #         paragraph.add_run(value)
    #     document.add_page_break()
    #     return document


    # labeling_df = pd.concat([
    #     table_df[table_df["project_relevant_idf_ratio"].between(1, 4)].sort_values("li_id")[:5],
    #     table_df[table_df["project_relevant_idf_ratio"].between(0.8, 1)].sort_values("li_id")[:5],
    #     table_df[table_df["project_relevant_idf_ratio"].between(0.6, 0.8)].sort_values("li_id")[:10],
    #     table_df[table_df["project_relevant_idf_ratio"].between(0.4, 0.6)].sort_values("li_id")[:10],
    #     table_df[table_df["project_relevant_idf_ratio"].between(-2, 0.4)].sort_values("li_id")[:5],
    # ])

    # document = docx.Document()

    # for row in labeling_df.sort_values("li_id").itertuples():
    #     document = add_founder(document, row)

    # document.save("/Users/zeintawil/Downloads/profiles_for_labeling.docx")
    return


@app.cell
def __():
    from collections import Counter
    return Counter,


@app.cell
def __(Counter, defaultdict, founders_df):
    top_tier_titles = defaultdict(int)
    _collections = founders_df[founders_df[("project_relevant_idf_ratio", "median")].between(1, 4)].sort_values("li_id")["member_experience_collection"].tolist()
    for _collection in _collections:
        for _experience in _collection:
            top_tier_titles[_experience["title"]] += 1
    Counter(top_tier_titles).most_common(50)
    return top_tier_titles,


@app.cell
def __(Counter, defaultdict, founders_df):
    low_tier_titles = defaultdict(int)
    _collections = founders_df[founders_df[("project_relevant_idf_ratio", "median")].between(-2, .4)].sort_values("li_id")["member_experience_collection"].tolist()
    for _collection in _collections:
        for _experience in _collection:
            low_tier_titles[_experience["title"]] += 1
    Counter(low_tier_titles).most_common(50)
    return low_tier_titles,


@app.cell
def __(metrics_df, px):
    px.histogram(metrics_df[("project_relevant_idf_ratio", "median")].values, histnorm="percent")
    return


@app.cell
def __(metrics_df):
    metrics_df[("project_relevant_idf_ratio", "median")].reset_index()
    return


@app.cell
def __():
    return


if __name__ == "__main__":
    app.run()
